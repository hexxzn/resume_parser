from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
import os.path
import pickle
import re

class Candidate:
    def __init__(self, first, last, location, email, phone, summary, skills, education, experience):
        # Candidate Identity
        self.first = first.contents[0] if first else 'First Name'
        self.last = last.contents[0] if last else 'Last Name'
        self.location = location.contents[0] if location else 'Location'
        self.email = email.contents[0] if email else 'Email Address'
        self.phone = phone.contents[0] if phone else 'Phone Number'
        self.summary = summary.contents[0] if summary else 'Summary'

        # Candidate Skills
        self.skills = '' if skills else 'Skills'
        for skill in skills:
            self.skills += format_string(skill.contents[0]) + '\n'
        self.skills = self.skills[:-1]

        # Candidate Education
        education_dict = {}
        if education:
            for item in education:
                degree_title = item.find('h3', {'data-shield-id': 'education_edu_title'})
                degree_title = degree_title.contents[0] if degree_title else 'Degree Title'
                university_name = item.find('span', {'data-shield-id': 'education_edu_school_span'})
                university_name = university_name.contents[0] if university_name else 'University Name'
                university_location = item.find('span', {'data-shield-id': 'education_edu_location_span'})
                university_location = university_location.contents[0] if university_location else 'University Location'
                attendance_date = item.find('div', {'data-shield-id': 'education_edu_dates'})
                attendance_date = attendance_date.contents[0] if attendance_date else 'Attendance Date'
                education_dict[degree_title] = university_name, university_location, attendance_date
        else:
            education_dict['Education'] = 'None'

        self.education = '' if education_dict else 'Education'
        for key in education_dict:
            self.education += format_string(key) + '\n'
            for value in education_dict[key]:
                self.education += format_string(value) + '\n'
            self.education += '\n'
        self.education = self.education[:-2]

        # Candidate Experience
        experience_dict = {}
        if experience:
            for item in experience:
                job_title = item.find('h3', {'data-shield-id': 'workExperience_work_title'})
                job_title = job_title.contents[0] if job_title else 'Job Title'
                company_name = item.find('span', {'data-shield-id': 'workExperience_work_experience_company'})
                company_name = company_name.contents[0] if company_name else 'Company Name'
                company_location = item.find('span', {'data-shield-id': 'workExperience_location_span'})
                company_location = company_location.contents[0] if company_location else 'Company Location'
                attendance_date = item.find('div', {'data-shield-id': 'workExperience_work_dates'})
                attendance_date = attendance_date.contents[0] if attendance_date else 'Attendance Date'
                job_description = item.find('p', {'data-shield-id': 'workExperience_work_description'})
                job_description = job_description.contents[0] if job_description else 'Job Description'
                experience_dict[job_title] = company_name, company_location, attendance_date, job_description
        else:
            experience_dict['Experience'] = 'None'

        self.experience = ''
        for key in experience_dict:
            self.experience += format_string(key) + '\n'
            for value in experience_dict[key]:
                self.experience += format_string(value) + '\n'
            self.experience += '\n'
        self.experience = self.experience[:-2]

# Extract Candidate Data From HTML
def main(url, manual_login):
    driver = webdriver.Firefox() # Create driver

    indeed_login(driver, url, manual_login) # Log in to Indeed
    
    source = driver.page_source # Get page source
    driver.close()
    
    candidate = create_candidate(source) # Create candidate

    parse_document_text(candidate) # Create resume from template

def indeed_login(driver, url, manual_login):
    #Log In Manually
    if not os.path.exists('resources\cookies.pkl') or manual_login:
        driver.get(url)
        WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.rdp-resume-container'))) # Wait for resume to load before saving cookies
        if EC.presence_of_element_located((By.CSS_SELECTOR, '.rdp-resume-container')):
            pickle.dump(driver.get_cookies(), open('resources/cookies.pkl', 'wb'))
    # Log In With Cookies
    else:
        driver.get(url)
        cookies = pickle.load(open('resources/cookies.pkl', 'rb'))
        for cookie in cookies:
            if cookie['domain'] == '.indeed.com':
                driver.add_cookie(cookie)
        driver.get(url)
        WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.rdp-resume-container')))

# Create Candidate Instance
def create_candidate(source):
    soup = BeautifulSoup(source, "html.parser")

    candidate = Candidate(
        soup.find('span', {'data-shield-id': 'firstname'}),
        soup.find('span', {'data-shield-id': 'lastname'}),
        soup.find('span', {'data-shield-id': 'locality'}),
        soup.find('div', {'data-shield-id': 'email'}),
        soup.find('div', {'data-shield-id': 'phone_number'}),
        soup.find('div', {'data-shield-id': 'res_summary'}),
        soup.findAll('span', {'data-shield-id': 'skill-text'}),
        soup.findAll('div', {'data-shield-id': 'education_data_display'}),
        soup.findAll('div', {'data-shield-id': 'workExperience_data_display'})
        )

    return candidate

# Iterate Through Text Elements In Document
def parse_document_text(candidate, document_name='resumes/template.docx'):
    document = Document(document_name)

    # Replace Text In Paragraphs
    for paragraph in document.paragraphs:
        replace_document_text(paragraph, candidate)

    # Replace Text In Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_document_text(paragraph, candidate)

    # Save Resume As New Word Document
    filename = candidate.first + candidate.last
    document.save('resumes/' + filename + '.docx')

# Find And Replace Text In Word Document
def replace_document_text(element, candidate):
    # String To Find, Replacement String
    replacements = [
        ['FirstName', candidate.first.upper()],
        ['LastName', candidate.last.upper()],
        ['PhoneNumber', candidate.phone],
        ['EmailAddress', candidate.email],
        ['Location', candidate.location],
        ['Summary', candidate.summary],
        ['Skills', candidate.skills],
        ['Education', candidate.education],
        ['Experience', candidate.experience],
        ]

    # Find Text To Replace
    inline = element.runs
    for i in range(len(inline)):
        for item in replacements:
            if item[0] in inline[i].text:
                item.append(i)

    # Replace Text Found
    for replacement in replacements:
        if len(replacement) == 3:
            index = replacement[2]
            target_text = replacement[0]
            replacement_text = replacement[1]
            inline[index].text = inline[index].text.replace(target_text, replacement_text)

# Remove Excess White Space And Line Breaks From String
def format_string(string):
    string = re.sub(' +', ' ', string)
    string = re.sub('\n', '', string)
    return string