from extract import extract
from replace import replace, format
from docx import Document

def insert(email, password, extract_url, document_name='template.docx', ):
    document = Document(document_name)
    resume_info = extract(email, password, extract_url)

    first_name = resume_info[0]
    last_name = resume_info[1]
    phone_number = resume_info[2]
    email_address = resume_info[3]
    location = resume_info[4]
    summary = resume_info[5]

    skills = ''
    for skill in resume_info[6]:
        skills += format(skill) + '\n'

    education = ''
    for key in resume_info[7]:
        education += format(key) + '\n'
        for value in resume_info[7][key]:
            education += format(value) + '\n'
        education += '\n'

    experience = ''
    for key in resume_info[8]:
        experience += format(key) + '\n'
        for value in resume_info[8][key]:
            experience += format(value) + '\n'
        experience += '\n'

    # Replace Text In Paragraphs
    for paragraph in document.paragraphs:
        replace(paragraph, 'FirstName', first_name.upper())
        replace(paragraph, 'LastName', last_name.upper())

    # Replace Text In Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace(paragraph, 'PhoneNumber', phone_number)
                    replace(paragraph, 'EmailAddress', email_address)
                    replace(paragraph, 'Location', location)
                    replace(paragraph, 'Summary', summary)
                    replace(paragraph, 'SkillsList', skills)
                    replace(paragraph, 'EducationList', education)
                    replace(paragraph, 'ExperienceList', experience)

    document.save('resume.docx')
