from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
import os.path
import pickle

from candidate import Candidate

# Extract candidate data from HTML.
def resume_extractor(url, manual_login):
    # Log in to Indeed and navigate to URL
    driver = webdriver.Firefox()
    indeed_login(driver, url, manual_login)

    # Get page source and close browser.
    source = driver.page_source
    driver.close()

    # Create candidate and insert candidate date into template.
    candidate = create_candidate(source)
    parse_document_text(candidate)

# Log in to Indeed.
def indeed_login(driver, url, manual_login):
    # Log in manually if manual login is checked or if cookies.pkl doesn't exist, otherwise auto login with cookies
    if not os.path.exists('resources\cookies.pkl') or manual_login:
        # Navigate to URL
        driver.get(url)

        # Wait for resume to fully load.
        WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.rdp-resume-container')))

        # If resume loads, save cookies.
        if EC.presence_of_element_located((By.CSS_SELECTOR, '.rdp-resume-container')):
            pickle.dump(driver.get_cookies(), open('resources/cookies.pkl', 'wb'))
    else:
        # Navigate to URL. (Lands on login page)
        driver.get(url)

        # Load cookies.
        cookies = pickle.load(open('resources/cookies.pkl', 'rb'))
        for cookie in cookies:
            if cookie['domain'] == '.indeed.com':
                driver.add_cookie(cookie)

        # Navigate to URL again after cookies loaded (Lands on resume page)
        driver.get(url)

        # Wait for resume to fully load.
        WebDriverWait(driver, 60).until(EC.presence_of_element_located((By.CSS_SELECTOR, '.rdp-resume-container')))

# Create candidate instance.
def create_candidate(source):
    # Load page source.
    soup = BeautifulSoup(source, "html.parser")

    # Create candidate with elements from page source.
    candidate = Candidate(
        soup.find('span', {'data-shield-id': 'firstname'}), # First name
        soup.find('span', {'data-shield-id': 'lastname'}), # Last name
        soup.find('span', {'data-shield-id': 'locality'}), # Location
        soup.find('div', {'data-shield-id': 'email'}), # Email address
        soup.find('div', {'data-shield-id': 'phone_number'}), # Phone number
        soup.find('div', {'data-shield-id': 'res_summary'}), # Summary
        soup.findAll('span', {'data-shield-id': 'skill-text'}), # Skills
        soup.findAll('div', {'data-shield-id': 'education_data_display'}), # Education
        soup.findAll('div', {'data-shield-id': 'workExperience_data_display'}) # Work experience
        )

    return candidate

# Iterate through text elements in document.
def parse_document_text(candidate, document='resumes/template.docx'):
    document = Document(document)

    # Replace text in paragraphs.
    for paragraph in document.paragraphs:
        replace_document_text(paragraph, candidate)

    # Replace text in tables.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_document_text(paragraph, candidate)

    # Save resume as new word document.
    filename = candidate.first + candidate.last
    document.save('resumes/' + filename + '.docx')

# Find and replace text in word document.
def replace_document_text(element, candidate):
    # [[string to replace, string to insert]]
    replacements = [
        ['CandidateFirst', candidate.first.upper()],
        ['CandidateLast', candidate.last.upper()],
        ['CandidatePhone', candidate.phone],
        ['CandidateEmail', candidate.email],
        ['CandidateLocation', candidate.location],
        ['CandidateSummary', candidate.summary],
        ['CandidateSkills', candidate.skills],
        ['CandidateEducation', candidate.education],
        ['CandidateExperience', candidate.experience],
        ]

    # Find text in template
    inline = element.runs

    # For line in element.
    for i in range(len(inline)):
        for item in replacements:
            # If text to replace is found in line of element.
            if item[0] in inline[i].text:
                # Append location of text to replace
                item.append(i)

    # Replace text in template
    for replacement in replacements:
        # If target text was found, its location becomes the 3rd element in nested replacements list.
        if len(replacement) == 3:
            # Inline location of text to replace.
            index = replacement[2]

            # Text to remove/replace.
            target_text = replacement[0]

            # Text to add/insert.
            replacement_text = replacement[1]

            # Replace
            inline[index].text = inline[index].text.replace(target_text, replacement_text)